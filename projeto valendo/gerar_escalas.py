import json
import sys
import subprocess
from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QVBoxLayout, QLabel, QCheckBox, QPushButton, QCalendarWidget, QDialog, QTextEdit
from PyQt5.QtGui import QFont
from PyQt5.QtCore import Qt
from docx import Document
from fpdf import FPDF
from docx import Document

def gerar_escala_docx(data, lista_alunos, caminho_modelo='caminho/para/modelo.docx', caminho_saida='escala_gerada.docx'):
    # Carrega o modelo .docx
    doc = Document(caminho_modelo)

    # Substitui o campo de data no documento
    for paragrafo in doc.paragraphs:
        if '{{DATA}}' in paragrafo.text:
            paragrafo.text = paragrafo.text.replace('{{DATA}}', data)

    # Substitui cada instância de {{aluno}} com os nomes da lista de alunos
    aluno_idx = 0
    for paragrafo in doc.paragraphs:
        while '{{aluno}}' in paragrafo.text and aluno_idx < len(lista_alunos):
            paragrafo.text = paragrafo.text.replace('{{aluno}}', lista_alunos[aluno_idx], 1)
            aluno_idx += 1

    # Salva o documento com as alterações
    doc.save(caminho_saida)
    print(f"Escala gerada em: {caminho_saida}")


class GerarEscalas(QWidget):
    def __init__(self):
        super().__init__()
        self.data = self.load_data()
        self.used_students = self.load_used_students()
        self.selected_turmas = []
        self.start_date = None
        self.end_date = None
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Gerar Escalas')
        self.resize(600, 600)

        # Layout principal
        layout = QVBoxLayout()

        # Título
        title = QLabel('Gerar Escalas', self)
        title_font = QFont('Arial', 18, QFont.Bold)
        title.setFont(title_font)
        title.setAlignment(Qt.AlignCenter)
        layout.addWidget(title)

        # Turmas
        turmas_layout = QVBoxLayout()
        self.checkboxes = []
        for turma in self.data['turmas']:
            checkbox = QCheckBox(turma['nome'], self)
            checkbox.setFont(QFont('Arial', 12))
            self.checkboxes.append(checkbox)
            turmas_layout.addWidget(checkbox)
        layout.addLayout(turmas_layout)

        # Botão Definir Dias
        self.btnDefinirDias = QPushButton('Definir Dias', self)
        self.btnDefinirDias.setFont(QFont('Arial', 12))
        self.btnDefinirDias.clicked.connect(self.abrirCalendario)
        layout.addWidget(self.btnDefinirDias)

        # Botão Gerar Escala
        self.btnGerarEscala = QPushButton('Gerar Escala', self)
        self.btnGerarEscala.setFont(QFont('Arial', 12))
        self.btnGerarEscala.clicked.connect(self.gerarEscala)
        layout.addWidget(self.btnGerarEscala)

        # Botão Visualizar Escala
        self.btnVisualizar = QPushButton('Visualizar Escala', self)
        self.btnVisualizar.setFont(QFont('Arial', 12))
        self.btnVisualizar.clicked.connect(self.visualizarPDF)
        layout.addWidget(self.btnVisualizar)

        # Botão Salvar
        self.btnSalvar = QPushButton('Salvar', self)
        self.btnSalvar.setFont(QFont('Arial', 12))
        self.btnSalvar.clicked.connect(self.salvarEscala)
        layout.addWidget(self.btnSalvar)

        # Botão Cancelar
        self.btnCancelar = QPushButton('Cancelar', self)
        self.btnCancelar.setFont(QFont('Arial', 12))
        self.btnCancelar.clicked.connect(self.close)
        layout.addWidget(self.btnCancelar)

        self.setLayout(layout)

    def abrirCalendario(self):
        self.calendarDialog = CalendarDialog()
        if self.calendarDialog.exec_():
            self.start_date, self.end_date = self.calendarDialog.getSelectedDates()

    def gerarEscala(self):
        self.selected_turmas = [checkbox.text() for checkbox in self.checkboxes if checkbox.isChecked()]
        if not self.selected_turmas or not self.start_date or not self.end_date:
            return

        escala_texto = ""
        self.escalas = []
        for turma_nome in self.selected_turmas:
            for turma in self.data['turmas']:
                if turma['nome'] == turma_nome:
                    escala_texto += f"Turma: {turma_nome}\n"
                    
                    alunos = [aluno for aluno in sorted(turma['alunos']) if aluno not in self.used_students[turma_nome]]
                    chefe = alunos[0] if alunos else ''
                    subchefe = alunos[1] if len(alunos) > 1 else ''
                    limpeza = alunos[-3:][::-1] if len(alunos) >= 3 else [''] * (3 - len(alunos)) + alunos[::-1]

                    self.used_students[turma_nome].extend([chefe, subchefe] + limpeza)

                    escala_texto += f"Chefe de Turma: {chefe}\n"
                    escala_texto += f"Subchefe de Turma: {subchefe}\n"
                    escala_texto += "Alunos da Limpeza:\n"
                    for aluno in limpeza:
                        escala_texto += f"  {aluno}\n"
                    escala_texto += "\n"

                    # Adiciona as datas selecionadas
                    escala_texto += "Dias Selecionados:\n"
                    escala_texto += f"De: {self.start_date.toString('dd/MM/yyyy')}\n"
                    escala_texto += f"Até: {self.end_date.toString('dd/MM/yyyy')}\n"
                    escala_texto += "\n"

                    self.escalas.append({
                        'turma': turma_nome,
                        'chefe': chefe,
                        'subchefe': subchefe,
                        'limpeza': limpeza,
                        'start_date': self.start_date.toString('dd/MM/yyyy'),
                        'end_date': self.end_date.toString('dd/MM/yyyy')
                    })

        self.mostrarEscala(escala_texto)
        self.save_used_students()

    def mostrarEscala(self, texto):
        dialog = QDialog(self)
        dialog.setWindowTitle('Escala Gerada')
        dialog.resize(400, 300)

        layout = QVBoxLayout()

        text_edit = QTextEdit(dialog)
        text_edit.setFont(QFont('Arial', 12))
        text_edit.setText(texto)
        text_edit.setReadOnly(True)
        layout.addWidget(text_edit)

        btnOk = QPushButton('OK', dialog)
        btnOk.setFont(QFont('Arial', 12))
        btnOk.clicked.connect(dialog.accept)
        layout.addWidget(btnOk)

        dialog.setLayout(layout)
        dialog.exec_()

    def visualizarPDF(self):
        try:
            subprocess.Popen(['start', 'escala.pdf'], shell=True)  # Para Windows, use 'start'
        except Exception as e:
            print(f"Erro ao abrir o PDF: {e}")

    def salvarEscala(self):
        if not hasattr(self, 'escalas'):
            return
        try:
            with open('escalas.json', 'w', encoding='utf-8') as file:
                json.dump(self.escalas, file, ensure_ascii=False, indent=4)
            print("Escala salva com sucesso!")
        except Exception as e:
            print(f"Erro ao salvar o arquivo: {e}")

    def load_data(self):
        try:
            with open('data.json', 'r', encoding='utf-8') as file:
                return json.load(file)
        except FileNotFoundError:
            return {'turmas': []}

    def load_used_students(self):
        try:
            with open('used_students.json', 'r', encoding='utf-8') as file:
                return json.load(file)
        except FileNotFoundError:
            return {turma['nome']: [] for turma in self.data['turmas']}

    def save_used_students(self):
        try:
            with open('used_students.json', 'w', encoding='utf-8') as file:
                json.dump(self.used_students, file, ensure_ascii=False, indent=4)
            print("Status dos alunos salvo com sucesso!")
        except Exception as e:
            print(f"Erro ao salvar o arquivo: {e}")

class CalendarDialog(QDialog):
    def __init__(self):
        super().__init__()
        self.start_date = None
        self.end_date = None
        self.initUI()

    def initUI(self):
        self.setWindowTitle('Selecionar Intervalo de Dias')
        self.resize(400, 300)

        layout = QVBoxLayout()

        self.calendarStart = QCalendarWidget(self)
        self.calendarStart.setGridVisible(True)
        self.calendarStart.setFont(QFont('Arial', 12))
        layout.addWidget(self.calendarStart)

        self.calendarEnd = QCalendarWidget(self)
        self.calendarEnd.setGridVisible(True)
        self.calendarEnd.setFont(QFont('Arial', 12))
        layout.addWidget(self.calendarEnd)

        btnOk = QPushButton('OK', self)
        btnOk.setFont(QFont('Arial', 12))
        btnOk.clicked.connect(self.accept)
        layout.addWidget(btnOk)

        self.setLayout(layout)

    def accept(self):
        self.start_date = self.calendarStart.selectedDate()
        self.end_date = self.calendarEnd.selectedDate()
        super().accept()

    def getSelectedDates(self):
        return self.start_date, self.end_date

if __name__ == '__main__':
    app = QApplication(sys.argv)
    mainWindow = QMainWindow()
    mainWindow.setCentralWidget(GerarEscalas())
    mainWindow.show()
    sys.exit(app.exec_())
